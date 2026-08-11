from __future__ import annotations

import asyncio
import hashlib
import io
import json
import logging
import os
import re
import tempfile
import time
import unicodedata
import uuid
from collections.abc import Awaitable, Callable, Sequence
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Protocol

import tiktoken
from docx import Document as DocxDocument
from pydantic import ValidationError
from pypdf import PdfReader

from ..core.config import Settings
from ..core.diagnostics import log_event
from ..core.errors import AppError, ConfigurationAppError, ProviderAppError, ValidationAppError
from ..models import (
    Chunk,
    DocumentSection,
    IngestionFileResult,
    IngestionResponse,
    SourceDocument,
)

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
METADATA_LINE = re.compile(r"^([A-Z][A-Z0-9_]+):\s*(.+?)\s*$")
HEADING = re.compile(r"^(?:#{1,6}\s+)?(?:\d+\.\s+)?[A-Z][A-Z0-9 &/(),:'_-]{2,}$")
SENTENCE = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9#])")
PATIENT_ID = re.compile(r"^SYN-P\d{3}$")
DOCUMENT_ID = re.compile(r"^DOC-(?:CLIN-P\d{3}|GUIDE|POLICY)-\d{3}$")
PATIENT_REFERENCE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9 ._-]{1,79}$")
NOTICE_MARKERS = (
    "SYNTHETIC MOCK DOCUMENT",
    "not intended for clinical care",
)
LOGGER = logging.getLogger("healthchat.ingestion")
METADATA_ORDER = (
    "DOCUMENT_ID",
    "DOCUMENT_TITLE",
    "DOCUMENT_TYPE",
    "DATA_CLASSIFICATION",
    "DOCUMENT_VERSION",
    "PATIENT_ID",
    "ENCOUNTER_ID",
    "ENCOUNTER_DATE",
    "AUTHOR_ROLE",
    "ORGANIZATION",
    "DOCUMENT_OWNER",
    "EFFECTIVE_DATE",
)


@dataclass(frozen=True)
class UploadPayload:
    filename: str
    content: bytes
    content_type: str | None = None


@dataclass(frozen=True)
class IndexedBatch:
    changed_chunk_ids: set[str]
    unchanged_chunk_ids: set[str]


ProgressCallback = Callable[[str, int, str], Awaitable[None]]


class EmbeddingService(Protocol):
    deployment: str

    async def embed(self, texts: list[str]) -> tuple[list[list[float]], int]: ...


class SearchIndexService(Protocol):
    index_name: str

    async def existing_hashes(self, document_ids: list[str]) -> dict[str, str]: ...

    async def ensure_schema(self, vector_dimensions: int) -> None: ...

    async def replace_documents(
        self,
        documents: list[SourceDocument],
        chunks: list[Chunk],
        vectors: dict[str, list[float]],
        existing_hashes: dict[str, str],
    ) -> IndexedBatch: ...


def _normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKC", value).replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+\n", "\n", value)
    return re.sub(r"\n{3,}", "\n\n", value).strip()


def _read_pdf(data: bytes, filename: str) -> list[tuple[int, str]]:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValidationAppError(f"{filename}: invalid PDF ({type(exc).__name__})") from exc
    if reader.is_encrypted:
        raise ValidationAppError(f"{filename}: encrypted PDFs are not supported")
    pages: list[tuple[int, str]] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            raise ValidationAppError(
                f"{filename}: PDF text extraction failed on page {number}"
            ) from exc
        if text.strip():
            pages.append((number, _normalize_text(text)))
    if not pages:
        raise ValidationAppError(
            f"{filename}: no extractable text found; image-only PDFs require OCR "
            "and are unsupported"
        )
    return pages


def _read_text(data: bytes, filename: str) -> list[tuple[int | None, str]]:
    try:
        decoded = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValidationAppError(f"{filename}: expected UTF-8 text") from exc
    return [(None, _normalize_text(decoded))]


def _read_docx(data: bytes, filename: str) -> list[tuple[int | None, str]]:
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise ValidationAppError(f"{filename}: invalid DOCX ({type(exc).__name__})") from exc
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        lines.append(f"# {text}" if style.startswith("heading") else text)
    content = _normalize_text("\n\n".join(lines))
    if not content:
        raise ValidationAppError(f"{filename}: no extractable text found")
    return [(None, content)]


def _metadata_and_body(text: str, filename: str) -> tuple[dict[str, str], str]:
    lines = text.splitlines()
    metadata: dict[str, str] = {}
    body_start = 0
    if lines and lines[0].strip() == "---":
        try:
            end = next(
                index for index, line in enumerate(lines[1:], start=1) if line.strip() == "---"
            )
        except StopIteration as exc:
            raise ValidationAppError(f"{filename}: unterminated metadata front matter") from exc
        header_lines = lines[1:end]
        body_start = end + 1
    else:
        header_lines = []
        for index, line in enumerate(lines):
            if not line.strip():
                body_start = index + 1
                break
            header_lines.append(line)
        else:
            body_start = len(lines)
    for line in header_lines:
        match = METADATA_LINE.match(line.strip())
        if not match:
            raise ValidationAppError(f"{filename}: invalid metadata line {line!r}")
        key, value = match.groups()
        if key in metadata:
            raise ValidationAppError(f"{filename}: duplicate metadata field {key}")
        metadata[key] = value.strip()
    return metadata, "\n".join(lines[body_start:]).strip()


def _source_role(document_id: str) -> str:
    if document_id.startswith("DOC-CLIN-"):
        return "clinical_note"
    if document_id.startswith("DOC-GUIDE-"):
        return "documentation_guide"
    if document_id.startswith("DOC-POLICY-"):
        return "payer_policy"
    raise ValidationAppError(f"{document_id}: unsupported document ID prefix")


def _parse_date(value: str | None, field: str, filename: str) -> date | None:
    if not value:
        return None
    try:
        return date.fromisoformat(value)
    except ValueError as exc:
        raise ValidationAppError(f"{filename}: {field} must be an ISO-8601 date") from exc


def _extract_sections(
    body: str, page_lookup: list[tuple[int | None, str]]
) -> list[DocumentSection]:
    upper = body.upper()
    start = 0
    for marker in NOTICE_MARKERS:
        marker_at = upper.find(marker.upper())
        if marker_at >= 0:
            start = max(start, marker_at + len(marker))
    evidence = body[start:]
    review_at = evidence.upper().find("CURATION REVIEW NOTES")
    if review_at >= 0:
        evidence = evidence[:review_at]
    evidence = evidence.strip()

    sections: list[DocumentSection] = []
    current_heading: str | None = None
    content_lines: list[str] = []

    def flush() -> None:
        nonlocal content_lines
        content = _normalize_text("\n".join(content_lines))
        if current_heading and content:
            page_numbers = [
                page
                for page, page_text in page_lookup
                if page is not None and (current_heading in page_text or content[:80] in page_text)
            ]
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content=content,
                    page_start=min(page_numbers) if page_numbers else None,
                    page_end=max(page_numbers) if page_numbers else None,
                )
            )
        content_lines = []

    for raw_line in evidence.splitlines():
        line = raw_line.strip()
        cleaned_heading = re.sub(r"^#{1,6}\s+", "", line)
        if line.startswith("#") or HEADING.fullmatch(line) or HEADING.fullmatch(cleaned_heading):
            flush()
            current_heading = cleaned_heading
            continue
        if current_heading is not None:
            content_lines.append(raw_line.rstrip())
    flush()
    if sections:
        return sections
    if any(page is not None for page, _ in page_lookup):
        return [
            DocumentSection(
                heading=f"Page {page}",
                content=text,
                page_start=page,
                page_end=page,
            )
            for page, text in page_lookup
            if page is not None and text.strip()
        ]
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n", evidence) if item.strip()]
    return [
        DocumentSection(heading=f"Document section {index}", content=paragraph)
        for index, paragraph in enumerate(paragraphs, start=1)
    ]


def normalize_patient_reference(value: str) -> str:
    normalized = " ".join(value.strip().split())
    if not PATIENT_REFERENCE.fullmatch(normalized):
        raise ValidationAppError(
            "Patient reference must be 2-80 characters using letters, numbers, spaces, dots, "
            "underscores, or hyphens"
        )
    canonical = re.fullmatch(r"syn-p(\d{1,3})", normalized, re.IGNORECASE)
    if canonical:
        return f"SYN-P{int(canonical.group(1)):03d}"
    return normalized


def _generic_document_id(patient_reference: str, filename: str) -> str:
    identity = f"{patient_reference.casefold()}\n{Path(filename).name.casefold()}"
    return f"DOC-CHART-{hashlib.sha256(identity.encode()).hexdigest()[:16].upper()}"


def _generic_document(
    upload: UploadPayload,
    pages: list[tuple[int | None, str]],
    patient_reference: str,
) -> SourceDocument:
    raw = _normalize_text("\n\n".join(text for _, text in pages))
    sections = _extract_sections(raw, pages)
    if not sections:
        raise ValidationAppError(f"{upload.filename}: no readable chart content was found")
    title = Path(upload.filename).stem.replace("_", " ").replace("-", " ").strip()
    title = " ".join(title.split()) or "Patient chart"
    document_id = _generic_document_id(patient_reference, upload.filename)
    normalized_lines = [
        f"DOCUMENT_ID: {document_id}",
        f"DOCUMENT_TITLE: {title}",
        "DOCUMENT_TYPE: patient_chart",
        "DATA_CLASSIFICATION: SYNTHETIC",
        "DOCUMENT_VERSION: 1.0.0",
        f"PATIENT_ID: {patient_reference}",
    ]
    for section in sections:
        normalized_lines.extend(["", section.heading, section.content])
    return SourceDocument(
        document_id=document_id,
        title=title,
        document_type="patient_chart",
        data_classification="SYNTHETIC",
        document_version="1.0.0",
        source_role="clinical_note",
        patient_id=patient_reference,
        source_filename=Path(upload.filename).name,
        normalized_text="\n".join(normalized_lines).strip() + "\n",
        sections=sections,
    )


def parse_upload(
    upload: UploadPayload,
    *,
    patient_reference: str | None = None,
    synthetic_confirmed: bool = False,
) -> SourceDocument:
    extension = Path(upload.filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValidationAppError(
            f"{upload.filename}: unsupported extension; expected .txt, .md, .pdf, or .docx"
        )
    if extension == ".pdf":
        pages = _read_pdf(upload.content, upload.filename)
    elif extension == ".docx":
        pages = _read_docx(upload.content, upload.filename)
    else:
        pages = _read_text(upload.content, upload.filename)
    raw = _normalize_text("\n\n".join(text for _, text in pages))
    structured = bool(re.search(r"(?m)^DOCUMENT_ID:\s*", raw))
    if not structured:
        if not synthetic_confirmed:
            raise ValidationAppError(
                f"{upload.filename}: confirm that the uploaded chart is synthetic demo data"
            )
        if patient_reference is None:
            raise ValidationAppError(f"{upload.filename}: a patient reference is required")
        return _generic_document(
            upload,
            pages,
            normalize_patient_reference(patient_reference),
        )
    if not all(marker.lower() in raw.lower() for marker in NOTICE_MARKERS):
        raise ValidationAppError(f"{upload.filename}: required synthetic-data notice is missing")
    metadata, body = _metadata_and_body(raw, upload.filename)
    required = {
        "DOCUMENT_ID",
        "DOCUMENT_TITLE",
        "DOCUMENT_TYPE",
        "DATA_CLASSIFICATION",
        "DOCUMENT_VERSION",
    }
    missing = sorted(required - metadata.keys())
    if missing:
        raise ValidationAppError(f"{upload.filename}: missing metadata {missing}")
    if metadata["DATA_CLASSIFICATION"].upper() != "SYNTHETIC":
        raise ValidationAppError(f"{upload.filename}: only explicitly SYNTHETIC data is accepted")
    document_id = metadata["DOCUMENT_ID"]
    if not DOCUMENT_ID.fullmatch(document_id):
        raise ValidationAppError(f"{upload.filename}: invalid DOCUMENT_ID {document_id!r}")
    source_role = _source_role(document_id)
    if source_role == "clinical_note":
        clinical_required = {
            "PATIENT_ID",
            "ENCOUNTER_ID",
            "ENCOUNTER_DATE",
            "AUTHOR_ROLE",
            "ORGANIZATION",
        }
        missing_clinical = sorted(clinical_required - metadata.keys())
        if missing_clinical:
            raise ValidationAppError(
                f"{upload.filename}: missing clinical metadata {missing_clinical}"
            )
        if not PATIENT_ID.fullmatch(metadata["PATIENT_ID"]):
            raise ValidationAppError(f"{upload.filename}: invalid PATIENT_ID")
        if patient_reference is not None:
            entered = normalize_patient_reference(patient_reference)
            if entered.casefold() != metadata["PATIENT_ID"].casefold():
                raise ValidationAppError(
                    f"{upload.filename}: embedded PATIENT_ID {metadata['PATIENT_ID']} does not "
                    f"match entered patient reference {entered}"
                )
    else:
        general_required = {"DOCUMENT_OWNER", "EFFECTIVE_DATE"}
        missing_general = sorted(general_required - metadata.keys())
        if missing_general:
            raise ValidationAppError(
                f"{upload.filename}: missing guide/policy metadata {missing_general}"
            )
    sections = _extract_sections(body, pages)
    if not sections:
        raise ValidationAppError(f"{upload.filename}: no evidentiary sections were found")
    normalized_lines = [f"{key}: {metadata[key]}" for key in METADATA_ORDER if metadata.get(key)]
    for section in sections:
        normalized_lines.extend(["", section.heading, section.content])
    try:
        return SourceDocument(
            document_id=document_id,
            title=metadata["DOCUMENT_TITLE"],
            document_type=metadata["DOCUMENT_TYPE"],
            data_classification="SYNTHETIC",
            document_version=metadata["DOCUMENT_VERSION"],
            source_role=source_role,
            patient_id=metadata.get("PATIENT_ID"),
            encounter_id=metadata.get("ENCOUNTER_ID"),
            encounter_date=_parse_date(
                metadata.get("ENCOUNTER_DATE"), "ENCOUNTER_DATE", upload.filename
            ),
            author_role=metadata.get("AUTHOR_ROLE"),
            organization=metadata.get("ORGANIZATION"),
            document_owner=metadata.get("DOCUMENT_OWNER"),
            effective_date=_parse_date(
                metadata.get("EFFECTIVE_DATE"), "EFFECTIVE_DATE", upload.filename
            ),
            source_filename=upload.filename,
            normalized_text="\n".join(normalized_lines).strip() + "\n",
            sections=sections,
        )
    except ValidationError as exc:
        raise ValidationAppError(f"{upload.filename}: metadata validation failed: {exc}") from exc


class TokenCounter:
    def __init__(self) -> None:
        self.encoding = tiktoken.get_encoding("cl100k_base")

    def count(self, text: str) -> int:
        return len(self.encoding.encode(text))

    def tail(self, text: str, tokens: int) -> str:
        encoded = self.encoding.encode(text)
        return self.encoding.decode(encoded[-tokens:]) if encoded else ""


def _split_section(
    section: DocumentSection,
    counter: TokenCounter,
    max_tokens: int,
    overlap_tokens: int,
) -> list[DocumentSection]:
    full = f"{section.heading}\n{section.content}"
    if counter.count(full) <= max_tokens:
        return [section]
    sentences = [item.strip() for item in SENTENCE.split(section.content) if item.strip()]
    if len(sentences) == 1:
        words = section.content.split()
        sentences = [" ".join(words[index : index + 40]) for index in range(0, len(words), 40)]
    pieces: list[DocumentSection] = []
    current = ""
    part = 1
    for sentence in sentences:
        proposed = f"{current} {sentence}".strip()
        if current and counter.count(f"{section.heading}\n{proposed}") > max_tokens:
            pieces.append(
                DocumentSection(
                    heading=f"{section.heading} (part {part})",
                    content=current,
                    page_start=section.page_start,
                    page_end=section.page_end,
                )
            )
            part += 1
            overlap = counter.tail(current, overlap_tokens)
            current = f"{overlap} {sentence}".strip()
        else:
            current = proposed
    if current:
        pieces.append(
            DocumentSection(
                heading=f"{section.heading} (part {part})",
                content=current,
                page_start=section.page_start,
                page_end=section.page_end,
            )
        )
    return pieces


def chunk_document(document: SourceDocument, settings: Settings) -> list[Chunk]:
    counter = TokenCounter()
    units = [
        piece
        for section in document.sections
        for piece in _split_section(
            section,
            counter,
            settings.max_chunk_tokens,
            settings.chunk_overlap_tokens,
        )
    ]
    grouped: list[list[DocumentSection]] = []
    current: list[DocumentSection] = []
    for unit in units:
        proposed = current + [unit]
        proposed_text = "\n\n".join(f"{item.heading}\n{item.content}" for item in proposed)
        if current and counter.count(proposed_text) > settings.target_chunk_tokens:
            grouped.append(current)
            current = [unit]
        else:
            current = proposed
    if current:
        grouped.append(current)

    chunks: list[Chunk] = []
    for order, group in enumerate(grouped, start=1):
        content = "\n\n".join(f"{item.heading}\n{item.content}" for item in group)
        content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
        pages = [page for item in group for page in (item.page_start, item.page_end) if page]
        chunks.append(
            Chunk(
                chunk_id=f"{document.document_id}-chunk-{order:03d}",
                document_id=document.document_id,
                title=document.title,
                source_filename=document.source_filename,
                source_role=document.source_role,
                document_type=document.document_type,
                document_version=document.document_version,
                patient_id=document.patient_id,
                encounter_id=document.encounter_id,
                encounter_date=document.encounter_date,
                author_role=document.author_role,
                organization=document.organization,
                document_owner=document.document_owner,
                effective_date=document.effective_date,
                section=" / ".join(item.heading for item in group),
                page_start=min(pages) if pages else None,
                page_end=max(pages) if pages else None,
                order=order,
                token_count=counter.count(content),
                content=content,
                content_hash=content_hash,
            )
        )
    return chunks


def embedding_text(chunk: Chunk) -> str:
    return (
        f"Title: {chunk.title}\nSource role: {chunk.source_role}\n"
        f"Section: {chunk.section}\n{chunk.content}"
    )


class ArtifactStore:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.normalized_dir = root / "normalized"
        self.chunks_dir = root / "chunks"
        self.manifest_path = root / "manifest.json"

    def _atomic_write(self, target: Path, content: str) -> None:
        target.parent.mkdir(parents=True, exist_ok=True)
        descriptor, temporary = tempfile.mkstemp(prefix=f".{target.name}.", dir=target.parent)
        try:
            with os.fdopen(descriptor, "w", encoding="utf-8") as handle:
                handle.write(content)
            os.replace(temporary, target)
        finally:
            if os.path.exists(temporary):
                os.unlink(temporary)

    def save(self, documents: Sequence[SourceDocument], chunks: Sequence[Chunk]) -> None:
        by_document: dict[str, list[Chunk]] = {}
        for chunk in chunks:
            by_document.setdefault(chunk.document_id, []).append(chunk)
        manifest = self.load_manifest()
        for document in documents:
            document_chunks = by_document.get(document.document_id, [])
            self._atomic_write(
                self.normalized_dir / f"{document.document_id}.txt", document.normalized_text
            )
            jsonl = "".join(
                json.dumps(chunk.model_dump(mode="json"), sort_keys=True) + "\n"
                for chunk in document_chunks
            )
            self._atomic_write(self.chunks_dir / f"{document.document_id}.jsonl", jsonl)
            manifest[document.document_id] = {
                "title": document.title,
                "source_role": document.source_role,
                "patient_id": document.patient_id,
                "source_filename": document.source_filename,
                "chunks": [
                    {"chunk_id": chunk.chunk_id, "content_hash": chunk.content_hash}
                    for chunk in document_chunks
                ],
            }
        self._atomic_write(
            self.manifest_path, json.dumps(manifest, indent=2, sort_keys=True) + "\n"
        )

    def load_manifest(self) -> dict[str, dict]:
        if not self.manifest_path.exists():
            return {}
        try:
            payload = json.loads(self.manifest_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return {}
        return payload if isinstance(payload, dict) else {}

    def load_chunks(self) -> list[Chunk]:
        chunks: list[Chunk] = []
        if not self.chunks_dir.exists():
            return chunks
        for path in sorted(self.chunks_dir.glob("*.jsonl")):
            for line in path.read_text(encoding="utf-8").splitlines():
                if line.strip():
                    chunks.append(Chunk.model_validate_json(line))
        return chunks


class IngestionService:
    def __init__(
        self,
        settings: Settings,
        embeddings: EmbeddingService | None,
        search: SearchIndexService | None,
        artifacts: ArtifactStore,
    ) -> None:
        self.settings = settings
        self.embeddings = embeddings
        self.search = search
        self.artifacts = artifacts

    def validate_uploads(
        self,
        uploads: Sequence[UploadPayload],
        *,
        patient_reference: str | None = None,
        synthetic_confirmed: bool = False,
    ) -> list[SourceDocument]:
        if not uploads:
            raise ValidationAppError("At least one file is required")
        if len(uploads) > self.settings.max_upload_files:
            raise ValidationAppError(f"At most {self.settings.max_upload_files} files are allowed")
        total = sum(len(upload.content) for upload in uploads)
        if total > self.settings.max_total_upload_bytes:
            raise ValidationAppError("Total upload size exceeds the configured limit")
        for upload in uploads:
            if len(upload.content) > self.settings.max_file_bytes:
                raise ValidationAppError(
                    f"{upload.filename}: file exceeds the configured size limit"
                )
        documents = [
            parse_upload(
                upload,
                patient_reference=patient_reference,
                synthetic_confirmed=synthetic_confirmed,
            )
            for upload in uploads
        ]
        ids = [document.document_id for document in documents]
        if len(ids) != len(set(ids)):
            raise ValidationAppError("The upload contains duplicate DOCUMENT_ID values")
        return documents

    async def ingest(
        self,
        uploads: Sequence[UploadPayload],
        *,
        patient_reference: str | None = None,
        synthetic_confirmed: bool = False,
        progress: ProgressCallback | None = None,
    ) -> IngestionResponse:
        started = time.perf_counter()
        ingestion_id = str(uuid.uuid4())
        stage = "validation"
        log_event(
            LOGGER,
            logging.INFO,
            "ingestion.started",
            ingestion_id=ingestion_id,
            upload_count=len(uploads),
            total_bytes=sum(len(upload.content) for upload in uploads),
        )
        try:
            if progress:
                await progress("validating", 12, "Checking files and patient reference")
            if patient_reference is not None:
                patient_reference = normalize_patient_reference(patient_reference)
            if progress:
                await progress("parsing", 24, "Reading chart text and document structure")
            documents = self.validate_uploads(
                uploads,
                patient_reference=patient_reference,
                synthetic_confirmed=synthetic_confirmed,
            )
            if progress:
                await progress("chunking", 38, "Preparing chart sections for search")
            chunks = [
                chunk
                for document in documents
                for chunk in chunk_document(document, self.settings)
            ]
            log_event(
                LOGGER,
                logging.INFO,
                "ingestion.validation.completed",
                ingestion_id=ingestion_id,
                document_count=len(documents),
                chunk_count=len(chunks),
                document_ids=[document.document_id for document in documents],
            )
            if self.embeddings is None or self.search is None:
                raise ConfigurationAppError(
                    "Azure OpenAI embeddings and Azure AI Search admin access are required "
                    "for ingestion"
                )

            stage = "existing_hash_lookup"
            if progress:
                await progress("checking_changes", 50, "Checking which chart sections changed")
            stage_started = time.perf_counter()
            log_event(
                LOGGER,
                logging.INFO,
                "ingestion.hash_lookup.started",
                ingestion_id=ingestion_id,
                index_name=self.search.index_name,
                document_count=len(documents),
            )
            existing = await self.search.existing_hashes(
                [document.document_id for document in documents]
            )
            log_event(
                LOGGER,
                logging.INFO,
                "ingestion.hash_lookup.completed",
                ingestion_id=ingestion_id,
                existing_chunk_count=len(existing),
                latency_ms=round((time.perf_counter() - stage_started) * 1000, 2),
            )
            changed = [
                chunk for chunk in chunks if existing.get(chunk.chunk_id) != chunk.content_hash
            ]
            log_event(
                LOGGER,
                logging.INFO,
                "ingestion.change_set.computed",
                ingestion_id=ingestion_id,
                changed_chunk_count=len(changed),
                unchanged_chunk_count=len(chunks) - len(changed),
            )
            vectors: dict[str, list[float]] = {}
            if changed:
                stage = "embedding"
                if progress:
                    await progress("embedding", 64, "Creating search representations")
                stage_started = time.perf_counter()
                log_event(
                    LOGGER,
                    logging.INFO,
                    "ingestion.embedding.started",
                    ingestion_id=ingestion_id,
                    deployment=self.embeddings.deployment,
                    chunk_count=len(changed),
                )
                embedded, input_tokens = await self.embeddings.embed(
                    [embedding_text(chunk) for chunk in changed]
                )
                if len(embedded) != len(changed) or not embedded or not embedded[0]:
                    raise ProviderAppError(
                        "Embedding response did not match the changed chunk batch"
                    )
                log_event(
                    LOGGER,
                    logging.INFO,
                    "ingestion.embedding.completed",
                    ingestion_id=ingestion_id,
                    vector_count=len(embedded),
                    vector_dimensions=len(embedded[0]),
                    input_tokens=input_tokens,
                    latency_ms=round((time.perf_counter() - stage_started) * 1000, 2),
                )

                stage = "schema_validation"
                stage_started = time.perf_counter()
                log_event(
                    LOGGER,
                    logging.INFO,
                    "ingestion.schema_validation.started",
                    ingestion_id=ingestion_id,
                    index_name=self.search.index_name,
                    vector_dimensions=len(embedded[0]),
                )
                await self.search.ensure_schema(len(embedded[0]))
                log_event(
                    LOGGER,
                    logging.INFO,
                    "ingestion.schema_validation.completed",
                    ingestion_id=ingestion_id,
                    latency_ms=round((time.perf_counter() - stage_started) * 1000, 2),
                )
                vectors = {
                    chunk.chunk_id: vector for chunk, vector in zip(changed, embedded, strict=True)
                }

            stage = "index_update"
            if progress:
                await progress("updating_index", 82, "Updating the knowledge base")
            stage_started = time.perf_counter()
            log_event(
                LOGGER,
                logging.INFO,
                "ingestion.index_update.started",
                ingestion_id=ingestion_id,
                changed_chunk_count=len(changed),
            )
            outcome = await self.search.replace_documents(documents, chunks, vectors, existing)
            log_event(
                LOGGER,
                logging.INFO,
                "ingestion.index_update.completed",
                ingestion_id=ingestion_id,
                changed_chunk_count=len(outcome.changed_chunk_ids),
                unchanged_chunk_count=len(outcome.unchanged_chunk_ids),
                latency_ms=round((time.perf_counter() - stage_started) * 1000, 2),
            )
        except AppError as exc:
            log_event(
                LOGGER,
                logging.WARNING if exc.status_code < 500 else logging.ERROR,
                "ingestion.failed",
                error=exc,
                include_traceback=exc.status_code >= 500,
                ingestion_id=ingestion_id,
                stage=stage,
                category=exc.category,
                status_code=exc.status_code,
                retryable=exc.retryable,
                latency_ms=round((time.perf_counter() - started) * 1000, 2),
            )
            raise
        except Exception as exc:
            log_event(
                LOGGER,
                logging.ERROR,
                "ingestion.failed",
                error=exc,
                include_traceback=True,
                ingestion_id=ingestion_id,
                stage=stage,
                latency_ms=round((time.perf_counter() - started) * 1000, 2),
            )
            raise ProviderAppError(
                f"Ingestion failed during {stage} while calling an Azure dependency "
                f"({type(exc).__name__})"
            ) from exc

        stage = "artifact_persistence"
        if progress:
            await progress("saving", 94, "Saving local chart metadata")
        stage_started = time.perf_counter()
        log_event(
            LOGGER,
            logging.INFO,
            "ingestion.artifact_persistence.started",
            ingestion_id=ingestion_id,
            artifact_dir=str(self.artifacts.root),
        )
        await asyncio.to_thread(self.artifacts.save, documents, chunks)
        log_event(
            LOGGER,
            logging.INFO,
            "ingestion.artifact_persistence.completed",
            ingestion_id=ingestion_id,
            latency_ms=round((time.perf_counter() - stage_started) * 1000, 2),
        )
        by_document = {document.document_id: [] for document in documents}
        for chunk in chunks:
            by_document[chunk.document_id].append(chunk)
        results: list[IngestionFileResult] = []
        indexed_documents = 0
        for document in documents:
            document_chunks = by_document[document.document_id]
            is_changed = any(
                chunk.chunk_id in outcome.changed_chunk_ids for chunk in document_chunks
            )
            indexed_documents += int(is_changed)
            results.append(
                IngestionFileResult(
                    filename=document.source_filename,
                    document_id=document.document_id,
                    status="indexed" if is_changed else "unchanged",
                    chunks=len(document_chunks),
                )
            )
        response = IngestionResponse(
            ingestion_id=ingestion_id,
            status="indexed" if outcome.changed_chunk_ids else "unchanged",
            documents_received=len(documents),
            documents_indexed=indexed_documents,
            documents_unchanged=len(documents) - indexed_documents,
            chunks_indexed=len(outcome.changed_chunk_ids),
            chunks_unchanged=len(outcome.unchanged_chunk_ids),
            index_name=self.search.index_name,
            embedding_deployment=self.embeddings.deployment,
            files=results,
            latency_ms=round((time.perf_counter() - started) * 1000, 2),
        )
        log_event(
            LOGGER,
            logging.INFO,
            "ingestion.completed",
            ingestion_id=ingestion_id,
            status=response.status,
            documents_received=response.documents_received,
            documents_indexed=response.documents_indexed,
            chunks_indexed=response.chunks_indexed,
            latency_ms=response.latency_ms,
        )
        if progress:
            await progress("completed", 100, "Charts are ready for questions")
        return response
