from __future__ import annotations

import io
import time

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.core.errors import ValidationAppError
from backend.main import create_app
from backend.services.ingestion import IngestionService, UploadPayload, parse_upload
from backend.services.patients import resolve_patient
from backend.services.retrieval import _odata_filter

from .conftest import FakeEmbeddingService, FakeSearchIndex


def test_generic_text_and_markdown_charts_use_entered_patient_reference():
    text = UploadPayload(
        "renal-follow-up.txt",
        b"ASSESSMENT\nChronic kidney disease stage 3b.\n\nPLAN\nRepeat labs in three months.",
    )
    markdown = UploadPayload(
        "medications.md",
        b"# Medications\n\nMetformin 500 mg twice daily.",
    )

    first = parse_upload(text, patient_reference="Demo Patient 6", synthetic_confirmed=True)
    repeated = parse_upload(text, patient_reference="Demo Patient 6", synthetic_confirmed=True)
    other_patient = parse_upload(
        text, patient_reference="Demo Patient 7", synthetic_confirmed=True
    )
    md_document = parse_upload(
        markdown, patient_reference="Demo Patient 6", synthetic_confirmed=True
    )

    assert first.document_id == repeated.document_id
    assert first.document_id != other_patient.document_id
    assert first.patient_id == "Demo Patient 6"
    assert first.source_role == "clinical_note"
    assert md_document.sections[0].heading == "Medications"


def test_heading_free_chart_gets_a_fallback_section():
    document = parse_upload(
        UploadPayload("plain.txt", b"The patient reports improved breathing after treatment."),
        patient_reference="CHART-42",
        synthetic_confirmed=True,
    )

    assert document.sections[0].heading == "Document section 1"
    assert "improved breathing" in document.sections[0].content


def test_docx_chart_preserves_heading_structure():
    stream = io.BytesIO()
    document = Document()
    document.add_heading("Assessment", level=1)
    document.add_paragraph("Asthma is documented as active.")
    document.add_heading("Plan", level=1)
    document.add_paragraph("Continue the current inhaler.")
    document.save(stream)

    parsed = parse_upload(
        UploadPayload(
            "pulmonary.docx",
            stream.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        patient_reference="SYN-P006",
        synthetic_confirmed=True,
    )

    assert [section.heading for section in parsed.sections] == ["Assessment", "Plan"]


def test_generic_upload_requires_confirmation_and_valid_reference():
    upload = UploadPayload("chart.txt", b"Assessment\nStable hypertension.")
    with pytest.raises(ValidationAppError, match="confirm"):
        parse_upload(upload, patient_reference="SYN-P006")
    with pytest.raises(ValidationAppError, match="Patient reference"):
        parse_upload(upload, patient_reference="?", synthetic_confirmed=True)


def test_embedded_patient_id_must_match_entered_reference(curated_uploads):
    with pytest.raises(ValidationAppError, match="does not match"):
        parse_upload(
            curated_uploads[0],
            patient_reference="SYN-P999",
            synthetic_confirmed=True,
        )


def test_duplicate_generic_filenames_are_rejected(settings, artifacts):
    service = IngestionService(settings, None, None, artifacts)
    uploads = [
        UploadPayload("visit.txt", b"Assessment\nAsthma."),
        UploadPayload("visit.txt", b"Assessment\nHypertension."),
    ]
    with pytest.raises(ValidationAppError, match="duplicate DOCUMENT_ID"):
        service.validate_uploads(
            uploads,
            patient_reference="SYN-P006",
            synthetic_confirmed=True,
        )


@pytest.mark.parametrize(
    ("query", "expected"),
    [
        ("What is documented for SYN-P004?", "SYN-P004"),
        ("What is documented for patient 4?", "SYN-P004"),
        ("What is documented for Demo Patient Six?", "Demo Patient Six"),
        ("What is documented for chart 42?", "CHART-42"),
    ],
)
def test_natural_patient_resolution(query, expected):
    references = ["SYN-P001", "SYN-P004", "Demo Patient Six", "CHART-42"]
    resolution = resolve_patient(query, references)
    assert resolution.status == "resolved"
    assert resolution.reference == expected


def test_patient_resolution_never_guesses_missing_ambiguous_or_unknown():
    references = ["SYN-P004", "MEMBER-004", "SYN-P005"]
    assert resolve_patient("Summarize the chart", references).status == "missing"
    assert resolve_patient("Summarize patient 4", references).status == "ambiguous"
    assert resolve_patient("Summarize SYN-P999", references).status == "unknown"


def test_resolved_name_is_used_as_an_exact_search_filter():
    assert _odata_filter("What medications are listed?", "Demo Patient 6") == (
        "(source_role eq 'clinical_note' and (patient_id eq 'Demo Patient 6'))"
    )


def test_background_ingestion_job_completes_and_refreshes_catalog(settings, artifacts):
    app = create_app(
        settings=settings,
        artifacts=artifacts,
        embeddings=FakeEmbeddingService(),
        search_admin=FakeSearchIndex(),
    )
    with TestClient(app) as client:
        created = client.post(
            "/ingestion/jobs",
            data={"patient_reference": "Demo Patient 6", "synthetic_confirmed": "true"},
            files=[
                (
                    "files",
                    (
                        "visit.txt",
                        b"ASSESSMENT\nAsthma is documented.\n\nPLAN\nContinue inhaler.",
                        "text/plain",
                    ),
                )
            ],
        )
        assert created.status_code == 202
        status_url = created.json()["status_url"]
        status = {}
        for _ in range(50):
            status = client.get(status_url).json()
            if status["status"] in {"completed", "failed"}:
                break
            time.sleep(0.01)
        catalog = client.get("/documents").json()

    assert status["status"] == "completed"
    assert status["stage"] == "completed"
    assert status["progress"] == 100
    assert status["result"]["documents_indexed"] == 1
    assert catalog[0]["patient_id"] == "Demo Patient 6"


def test_background_job_validation_failure_and_unknown_job(settings, artifacts):
    app = create_app(
        settings=settings,
        artifacts=artifacts,
        embeddings=FakeEmbeddingService(),
        search_admin=FakeSearchIndex(),
    )
    with TestClient(app) as client:
        missing = client.get("/ingestion/jobs/not-a-job")
        unconfirmed = client.post(
            "/ingestion/jobs",
            data={"patient_reference": "SYN-P006", "synthetic_confirmed": "false"},
            files=[("files", ("visit.txt", b"Assessment\nAsthma.", "text/plain"))],
        )
        created = client.post(
            "/ingestion/jobs",
            data={"patient_reference": "SYN-P006", "synthetic_confirmed": "true"},
            files=[("files", ("scan.csv", b"bad,file", "text/csv"))],
        )
        failed = {}
        for _ in range(50):
            failed = client.get(created.json()["status_url"]).json()
            if failed["status"] == "failed":
                break
            time.sleep(0.01)

    assert missing.status_code == 404
    assert unconfirmed.status_code == 422
    assert unconfirmed.json()["category"] == "validation"
    assert created.status_code == 202
    assert failed["stage"] == "failed"
    assert failed["error"]["category"] == "validation"
